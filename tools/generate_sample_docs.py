"""Generate the binary sample documents (PDF + DOCX) for the knowledge base.

Why a generator instead of committing binaries? Binary files can't be edited or
diffed easily. Keeping the *content* here as plain text means students can read
it, tweak it, and regenerate the documents — and it doubles as a worked example
of producing PDFs (fpdf2) and Word files (python-docx) programmatically.

Run it::

    pip install fpdf2 python-docx
    python tools/generate_sample_docs.py

It writes into data/<role>/ alongside the Markdown documents. The ingestion
pipeline reads all of them the same way and cites their section/page.
"""

from __future__ import annotations

import sys
from pathlib import Path

DATA = Path(__file__).resolve().parent.parent / "data"


# ---------------------------------------------------------------------------
# PDF helpers (fpdf2). Core PDF fonts are latin-1 only, so we keep PDF text
# ASCII-safe — the helper below downgrades common typographic characters.
# ---------------------------------------------------------------------------
def _ascii(text: str) -> str:
    repl = {"—": "-", "–": "-", "•": "-", "’": "'", "‘": "'",
            "“": '"', "”": '"', "→": "->", "±": "+/-", "⌖": "TP",
            "°": " deg", "≥": ">=", "≤": "<=", " ": " "}
    for bad, good in repl.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def make_pdf(path: Path, title: str, blocks: list[tuple[str, str]]) -> None:
    """Render a simple titled PDF.

    ``blocks`` is a list of (kind, text) where kind is 'h' (heading),
    'p' (paragraph), or 'li' (bullet item).
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Return to the left margin and advance after each block, otherwise fpdf2
    # leaves the cursor at the right edge (zero usable width -> crash).
    def line(h: int, txt: str) -> None:
        pdf.multi_cell(0, h, _ascii(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "B", 16)
    line(10, title)
    pdf.ln(2)

    for kind, text in blocks:
        if kind == "h":
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 13)
            line(8, text)
        elif kind == "li":
            pdf.set_font("Helvetica", "", 11)
            line(6, f"  -  {text}")
        else:
            pdf.set_font("Helvetica", "", 11)
            line(6, text)
    pdf.output(str(path))
    print(f"  wrote {path.relative_to(DATA.parent)}")


# ---------------------------------------------------------------------------
# DOCX helpers (python-docx). Word handles Unicode natively.
# ---------------------------------------------------------------------------
def make_docx(path: Path, title: str, builder) -> None:
    import docx

    document = docx.Document()
    document.add_heading(title, level=0)
    builder(document)
    document.save(str(path))
    print(f"  wrote {path.relative_to(DATA.parent)}")


# ===========================================================================
# PDF content
# ===========================================================================
def maintenance_pm_schedule_pdf() -> None:
    make_pdf(
        DATA / "maintenance" / "preventive_maintenance_schedule.pdf",
        "Preventive Maintenance Schedule - Line 2",
        [
            ("p", "Preventive maintenance (PM) keeps equipment reliable and prevents "
                  "the breakdowns that cause unplanned downtime. Follow this schedule; "
                  "record completion in the maintenance log."),
            ("h", "VMX-500 CNC Mill"),
            ("li", "Daily: check coolant level, air pressure 6 bar, run 8-min warm-up."),
            ("li", "Weekly: check way-oil reservoir, coolant concentration 6-8%, clean chip tray and filter."),
            ("li", "Every 500 hours: grease X/Y/Z ballscrew bearings (NLGI 2), replace coolant filter."),
            ("li", "Every 3 months: replace coolant, inspect spindle drive belt tension."),
            ("h", "BX-200 Conveyor"),
            ("li", "Daily: check belt tracking, listen for abnormal noise, confirm guards."),
            ("li", "Weekly: inspect rollers, check take-up tension, clear debris."),
            ("li", "Monthly: gearmotor oil level, belt wear, bolt torque."),
            ("li", "Every 1,000 hours: replace worn idlers, inspect drive lagging."),
            ("h", "Rules"),
            ("li", "Apply lockout/tagout (LOTO-01) before any task inside a guard."),
            ("li", "Use only calibrated tools and OEM spare parts."),
            ("li", "Log every PM with date, technician, and downtime."),
            ("h", "Escalation"),
            ("p", "Overdue PM or a fault found during PM: raise a work order to the "
                  "maintenance supervisor (ext. 4500)."),
        ],
    )


def safety_emergency_procedures_pdf() -> None:
    make_pdf(
        DATA / "safety" / "emergency_procedures.pdf",
        "Emergency Procedures - Quick Guide",
        [
            ("p", "Keep this guide visible at every workstation. In any emergency, "
                  "protect people first, then property. Know your nearest exit and "
                  "assembly point."),
            ("h", "Emergency stop"),
            ("li", "Press the nearest red E-stop to halt machinery immediately."),
            ("li", "Do not reset an E-stop until the hazard is cleared and the line is confirmed safe."),
            ("h", "Fire"),
            ("li", "Raise the alarm and call emergency services."),
            ("li", "Only fight a small fire with the correct extinguisher if trained and safe to do so."),
            ("li", "Evacuate via the nearest marked exit; do not use elevators."),
            ("li", "Go to the assembly point in the north car park; await the roll call."),
            ("h", "Injury / first aid"),
            ("li", "Make the area safe (LOTO machinery if needed) before approaching the casualty."),
            ("li", "Call a trained first aider; do not move a seriously injured person unless in danger."),
            ("li", "Report every injury and complete an incident report (see incident_report_form)."),
            ("h", "Chemical spill"),
            ("li", "Small spill: use the spill kit and PPE per the SDS (see chemical_safety)."),
            ("li", "Large spill or any spill reaching a drain: evacuate the area and call the Safety Officer."),
            ("h", "Key contacts"),
            ("p", "Safety Officer ext. 4911. Emergency services: dial the site emergency number "
                  "posted at each phone."),
        ],
    )


def quality_inspection_checklist_pdf() -> None:
    make_pdf(
        DATA / "quality" / "inspection_checklist.pdf",
        "First-Article Inspection Checklist",
        [
            ("p", "Complete this checklist on the first part after every setup or "
                  "changeover, BEFORE running the batch. Production starts only when "
                  "the first article passes and is signed off."),
            ("h", "Setup verification"),
            ("li", "Correct NC program number loaded and matches the drawing."),
            ("li", "Correct fixture installed; clamps torqued to the setup sheet."),
            ("li", "Work offset (G54) verified with edge finder or probe."),
            ("li", "All gauges present and in calibration (check sticker dates)."),
            ("h", "Dimensional checks (example part P-Bracket)"),
            ("li", "Overall length 120.0 mm, tolerance +/-0.10 mm - caliper CAL-07."),
            ("li", "Bore diameter 25.00 mm, +0.02 / -0.00 mm - bore gauge CAL-12."),
            ("li", "Surface finish Ra 1.6 max - profilometer."),
            ("li", "Hole position true-position 0.1 mm - CMM."),
            ("h", "Disposition"),
            ("li", "All features in tolerance -> sign off, release the batch."),
            ("li", "Any feature out of tolerance -> quarantine, raise an NCR, call Quality."),
            ("h", "Sign-off"),
            ("p", "Operator: ____________  Inspector: ____________  Date/Time: __________"),
            ("p", "Out-of-tolerance or missing spec? Do not guess - contact the Quality "
                  "Engineer (ext. 4700)."),
        ],
    )


# ===========================================================================
# DOCX content
# ===========================================================================
def maintenance_work_order_docx() -> None:
    def build(doc):
        doc.add_paragraph(
            "Raise a work order for any corrective or preventive maintenance task. "
            "Apply lockout/tagout (LOTO-01) before work inside any guard or enclosure."
        )
        doc.add_heading("Work Order Details", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for label in ["Work order no.", "Date raised", "Asset ID", "Raised by",
                      "Priority (Low/Med/High/Breakdown)", "Reported symptom / alarm code"]:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = ""
        doc.add_heading("Work Performed", level=1)
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Light Grid Accent 1"
        hdr = t2.rows[0].cells
        for i, h in enumerate(["Date", "Technician", "Action taken", "Parts used (P/N)"]):
            hdr[i].text = h
        for _ in range(4):
            t2.add_row()
        doc.add_heading("Close-out", level=1)
        doc.add_paragraph("Downtime (min): ________   Root cause: ____________________")
        doc.add_paragraph("LOTO applied (Y/N): ____   Returned to service by: ____________")
        doc.add_paragraph(
            "Log the completed work order in the maintenance log. Recurring faults "
            "must be escalated to the maintenance supervisor (ext. 4500)."
        )

    make_docx(DATA / "maintenance" / "work_order_template.docx", "Maintenance Work Order", build)


def safety_incident_report_docx() -> None:
    def build(doc):
        doc.add_paragraph(
            "Complete this form for every injury, near-miss, spill, or unsafe "
            "condition. Reporting is blameless and helps prevent recurrence. Submit "
            "to the Safety Officer (ext. 4911) within 24 hours."
        )
        doc.add_heading("Incident Details", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for label in ["Date & time", "Location / area", "Reported by", "People involved",
                      "Type (Injury / Near-miss / Spill / Other)", "Equipment involved"]:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = ""
        doc.add_heading("What happened", level=1)
        doc.add_paragraph("Describe the sequence of events, including any alarm or energy source:")
        for _ in range(3):
            doc.add_paragraph("_______________________________________________________________")
        doc.add_heading("Immediate action taken", level=1)
        doc.add_paragraph("(First aid, LOTO, spill containment, evacuation, etc.)")
        doc.add_paragraph("_______________________________________________________________")
        doc.add_heading("Follow-up", level=1)
        t2 = doc.add_table(rows=0, cols=2)
        t2.style = "Light Grid Accent 1"
        for label in ["Root cause", "Corrective action", "Owner", "Due date", "Status"]:
            row = t2.add_row().cells
            row[0].text = label
            row[1].text = ""
        doc.add_paragraph("Safety Officer sign-off: ____________________   Date: __________")

    make_docx(DATA / "safety" / "incident_report_form.docx", "Safety Incident Report", build)


def quality_ncr_docx() -> None:
    def build(doc):
        doc.add_paragraph(
            "Raise a Nonconformance Report (NCR) when product does not meet "
            "specification. Quarantine the affected parts in the red NCR bin. Only "
            "Quality may decide disposition - operators must not use-as-is on their own."
        )
        doc.add_heading("Nonconformance Details", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for label in ["NCR no.", "Date", "Part number / name", "Lot / batch",
                      "Quantity affected", "Raised by", "Operation / machine"]:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = ""
        doc.add_heading("Description of nonconformance", level=1)
        doc.add_paragraph("Feature / dimension out of spec, with measured vs required value:")
        t2 = doc.add_table(rows=1, cols=4)
        t2.style = "Light Grid Accent 1"
        hdr = t2.rows[0].cells
        for i, h in enumerate(["Feature", "Required (spec)", "Measured", "Gauge / CAL ID"]):
            hdr[i].text = h
        for _ in range(3):
            t2.add_row()
        doc.add_heading("Disposition (Quality use only)", level=1)
        doc.add_paragraph("[ ] Rework   [ ] Scrap   [ ] Use-as-is (concession)   [ ] Return to supplier")
        doc.add_paragraph("Root cause: ____________________   Corrective action: ____________________")
        doc.add_paragraph("Quality Engineer: ____________________   Date: __________")

    make_docx(DATA / "quality" / "nonconformance_report.docx", "Nonconformance Report (NCR)", build)


def main() -> int:
    print("Generating PDF documents...")
    maintenance_pm_schedule_pdf()
    safety_emergency_procedures_pdf()
    quality_inspection_checklist_pdf()
    print("Generating DOCX documents...")
    maintenance_work_order_docx()
    safety_incident_report_docx()
    quality_ncr_docx()
    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())

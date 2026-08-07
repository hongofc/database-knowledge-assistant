"""Load documents from disk and split them into retrievable chunks.

Supported formats: Markdown (``.md``), plain text (``.txt``), PDF (``.pdf``),
and Word (``.docx``) — a real factory knowledge base is full of PDFs and Word
SOPs, not just text. Text formats need nothing extra. PDF reading uses ``pypdf``
and Word reading uses ``python-docx``; both are *optional* — if a library is
missing those files are skipped with a one-time notice instead of crashing.

This module also captures **citation metadata** as it goes: the nearest section
heading and (for PDFs) the page number, so every retrieved chunk can point a
technician to the exact place in the official document.

To add another format (HTML, CAD export, ...), write a small reader that returns
plain text and register it in ``_READERS`` below — nothing else changes.
"""

from __future__ import annotations

import re
import warnings
from pathlib import Path

from .config import settings
from .rag.base import Chunk, Document

# Markers/patterns used to recover citation context during chunking.
_PAGE_MARKER = "<<<PAGE:{n}>>>"
_PAGE_RE = re.compile(r"^<<<PAGE:(\d+)>>>$")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _read_text_file(path: Path) -> str:
    """Markdown / plain text — just decode the bytes."""
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF, inserting page markers so we can cite pages."""
    try:
        from pypdf import PdfReader
    except ImportError:
        warnings.warn(
            "Skipping PDF files: install 'pypdf' to read them (pip install pypdf).",
            stacklevel=2,
        )
        return ""
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        parts.append(_PAGE_MARKER.format(n=i))
        parts.append(page.extract_text() or "")
    return "\n\n".join(parts)


def _read_docx(path: Path) -> str:
    """Extract paragraphs and tables; promote Word headings to '# ' lines.

    Promoting heading-styled paragraphs to Markdown headings lets the same
    section-tracking logic produce citations for Word documents too.
    """
    try:
        import docx  # python-docx
    except ImportError:
        warnings.warn(
            "Skipping .docx files: install 'python-docx' to read them "
            "(pip install python-docx).",
            stacklevel=2,
        )
        return ""
    document = docx.Document(str(path))
    parts: list[str] = []
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if style.startswith("heading") or style == "title":
            parts.append(f"# {text}")
        else:
            parts.append(text)
    # Flatten tables into 'cell | cell | cell' lines so they remain searchable.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


# Suffix -> reader. The single source of truth for "what can we ingest?".
_READERS = {
    ".md": _read_text_file,
    ".markdown": _read_text_file,
    ".txt": _read_text_file,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}

SUPPORTED_SUFFIXES = set(_READERS)


def load_documents(folder: Path, role: str) -> list[Document]:
    """Read every supported file under ``folder`` into a Document.

    ``role`` tags each document with the specialist it belongs to so chunks
    keep their provenance through the whole pipeline. Unreadable or unsupported
    files are skipped (with a notice for missing optional libraries) rather than
    aborting the whole ingest.
    """
    docs: list[Document] = []
    if not folder.exists():
        return docs
    for path in sorted(folder.rglob("*")):
        reader = _READERS.get(path.suffix.lower())
        if reader is None or not path.is_file():
            continue
        try:
            text = reader(path).strip()
        except Exception as exc:  # noqa: BLE001 - one bad file shouldn't stop ingest
            warnings.warn(f"Could not read {path.name}: {exc}", stacklevel=2)
            continue
        if not text:
            continue
        try:
            source = str(path.relative_to(_project_root()))
        except ValueError:
            source = str(path)
        docs.append(Document(text=text, source=source, role=role))
    return docs


def _project_root() -> Path:
    from .config import PROJECT_ROOT

    return PROJECT_ROOT


def chunk_document(
    doc: Document,
    chunk_size: int | None = None,
    overlap: int | None = None,
    strategy: str | None = None,
) -> list[Chunk]:
    """Split a document into chunks using the configured strategy.

    Dispatches to :mod:`factory_knowledge.chunking`, where four interchangeable
    strategies live (``fixed``, ``recursive``, ``semantic``, ``metadata_aware``).
    Set ``CHUNK_STRATEGY`` in ``.env`` or pass ``strategy=`` to override for a
    single call — which is exactly what the evaluation harness does to compare
    strategies on identical documents.

    The legacy character-window behaviour is preserved as ``fixed`` so it stays
    available as the control in those comparisons.
    """
    from .chunking import build_strategy

    kwargs: dict = {
        "chunk_size": chunk_size or settings.chunk_size,
        "overlap": overlap or settings.chunk_overlap,
    }
    name = strategy or settings.chunk_strategy
    if name == "semantic":
        kwargs["breakpoint_percentile"] = settings.breakpoint_percentile
    return build_strategy(name, **kwargs).split(doc)


def build_chunks(
    folder: Path, role: str, strategy: str | None = None
) -> list[Chunk]:
    """Convenience: load a folder and return all chunks for ``role``."""
    chunks: list[Chunk] = []
    for doc in load_documents(folder, role):
        chunks.extend(chunk_document(doc, strategy=strategy))
    return chunks
